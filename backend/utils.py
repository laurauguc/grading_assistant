import docx
import os
import json
import pathlib
import google.generativeai as genai
from dotenv import load_dotenv, find_dotenv
from pathlib import Path
import re
from openai import OpenAI

from . import schemas

class DocxToMarkdownConverter:
    """
    A class to convert DOCX files to Markdown format.

    This class provides methods to read a DOCX file, convert its content
    (including paragraphs and tables) to Markdown format, and save the
    converted content to a specified file path.
    """

    def __init__(self):
        """
        Initialize the converter with paths to the DOCX file and the output Markdown file.

        Parameters:
        docx_path (str): Path to the DOCX file.
        md_path (str): Path to the output Markdown file.
        """
        #self.docx_path = docx_path
        self.doc = None

    def read_word_file(self, docx_path):
        """
        Read the DOCX file and store the Document object in the instance variable.
        """
        self.doc = docx.Document(docx_path)

    def convert_paragraph_to_md(self, paragraph):
        """
        Convert a DOCX paragraph to a Markdown formatted string.

        Parameters:
        paragraph (docx.text.paragraph.Paragraph): The paragraph to convert.

        Returns:
        str: The Markdown formatted paragraph.
        """
        md_text = paragraph.text.strip()
        if paragraph.style.name.startswith('Heading'):
            try:
                level = int(paragraph.style.name.split()[-1])
            except:
                level = 1
            md_text = f"{'#' * level} {md_text}"
        elif paragraph.style.name == 'Normal' and md_text:
            md_text = f"{md_text}\n"
        return md_text

    def convert_cell_to_md(self, cell):
        """
        Convert the content of a DOCX table cell to a Markdown formatted string.

        Parameters:
        cell (docx.table._Cell): The cell to convert.

        Returns:
        str: The Markdown formatted cell content.
        """
        md_lines = []
        for paragraph in cell.paragraphs:
            if paragraph.text.strip():  # Only add non-empty paragraphs
                md_lines.append(paragraph.text.strip())
        for inner_table in cell.tables:
            for inner_row in inner_table.rows: # changed here
                for inner_cell in inner_row.cells:
                    for paragraph in inner_cell.paragraphs:
                        md_lines.append(paragraph.text.strip(". "))
        cell_md = '. '.join(md_lines).strip()  # Join paragraphs with spaces
        return cell_md

    def convert_table_to_md(self, table):
        """
        Convert a DOCX table to a Markdown formatted string.

        Parameters:
        table (docx.table.Table): The table to convert.

        Returns:
        list: A list of strings representing the Markdown formatted table.
        """
        md_lines = []
        header = True
        for row in table.rows:
            row_data = [self.convert_cell_to_md(cell) for cell in row.cells]
            #row_data = [cell.text.strip() for cell in row.cells]
            if len(set(row_data)) == 1: # wide row
                md_lines.append("\n"+row_data[0]+"\n")
                header = True
            else:
                md_line = '| ' + ' | '.join(row_data) + ' |'
                md_lines.append(md_line)
                if header:
                    separator = '| ' + ' | '.join(['---'] * len(row_data)) + ' |'
                    md_lines.append(separator)
                    header = False
        return md_lines

    def convert_docx_to_md(self):
        """
        Convert the content of the DOCX Document object to Markdown format.

        Returns:
        str: The entire document content converted to Markdown format.
        """
        md_lines = []
        for element in self.doc.element.body:
            if isinstance(element, docx.oxml.CT_P): # paragraph
                paragraph = docx.text.paragraph.Paragraph(element, self.doc)
                md_line = self.convert_paragraph_to_md(paragraph)
                if md_line:
                    md_lines.append(md_line)
            elif isinstance(element, docx.oxml.CT_Tbl): # table
                table = docx.table.Table(element, self.doc)
                md_table_lines = self.convert_table_to_md(table)
                md_lines.extend(md_table_lines)
            elif isinstance(element, docx.oxml.CT_SectPr):
                md_lines.append("<Image found here>") # TO DO: Replace <Image found here> with descriptive image metadata, OR: make an API call to obtain a description (Premium version).
        return "\n".join(md_lines)

    def convert(self, docx_path):
        """
        Orchestrate the conversion of a DOCX file to Markdown format and save the result to a file.
        """
        if not os.path.exists(docx_path):
            raise FileNotFoundError(f"The file {docx_path} does not exist.")

        self.read_word_file(docx_path)
        md_content = self.convert_docx_to_md()
        return md_content  # Return the Markdown content to the console

def format_suggestions_and_rewrite_prompt(grading_feedback, student_assignment):
    return f"""
    Based on the following grading feedback,

    1. Provide suggestions for improvement
    2. Rewrite the essay using the suggested improvements

    Grading Feedback:
    {grading_feedback}

    Student Assignment:
    {student_assignment}

    Take each improvement provided and identify,
    a. the smallest chunk of the original text where it was applied
    b. the criterion from the rubric. If the criterion is not available, use the category or create one.
    c. the reason for the suggested improvement

    Provide steps a, b, and c in the following JSON format:
    [
        {{"improvement": "improvement_1", "criterion_from_rubric": "criterion_from_rubric_1", "reason_for_suggestion": "reason_for_suggestion_1", "original_text": "original_text_1", "revised_text": "revised_text_1"}},
        ...
        {{"improvement": "improvement_n", "criterion_from_rubric": "criterion_from_rubric_n", "reason_for_suggestion": "reason_for_suggestion_n", "original_text": "original_text_n", "revised_text": "original_text_n"}}
    ]
    """

# def generate_suggestions_and_rewrite_essay(model, grading_feedback, student_assignment):
#     prompt = format_suggestions_and_rewrite_prompt(grading_feedback, student_assignment)
#
#     response = model.generate_content(prompt)
#
#     # Ensure the response streaming completes
#     response.resolve()
#
#     if response and response.text:
#         #print("Full response from model:\n", response.text)  # Debugging statement
#         return response.text
#     else:
#         #print("Error: No response from model")
#         return None

def get_feeback(system_prompt, user_prompt, response_format):
    client = OpenAI()
    response = client.chat.completions.create(
      model="gpt-4o",
        messages=[
            {
                "role": "system",
                "content": system_prompt,
            }, {
                "role": "user",
                "content": user_prompt
            }
        ],
      response_format = response_format,
      temperature=0
    )
    message_content = response.choices[0].message.content
    return json.loads(message_content)

def generate_feedback(grading_feedback, student_assignment):
    system_prompt_improvements = (
        "Role: You are a Teaching Assistant tasked with providing detailed and actionable feedback to students on their writing assignments."
        "\n\n"
        "The user will provide you with the student's assignment text and high-level feedback, both enclosed within triple backticks."
        "Your responsibility is to generate specific suggestions for improving the assignment based on the provided feedback."
        "\n\n"
        "For each suggested improvement, include the following elements:"
        "1. **improvement**: A concise description of the suggested change."
        "2. **original_text**: The exact excerpt from the student's assignment where the improvement applies. Ensure this matches a verbatim portion of the input text."
        "3. **criterion_from_rubric**: The related grading rubric criterion. If no explicit criterion is provided, derive an appropriate one based on the feedback."
        "4. **reason_for_suggestion**: A clear explanation of why the improvement is necessary or beneficial."
        "5. **examples**: Provide 1 or 2 example ideas for the suggested improvement."
        "\n\n"
        "Ensure that the `original_text` aligns with an exact excerpt from the student's assignment. Do not paraphrase or generalize this text."
    )

    system_prompt_strengths = (
        "Role: You are a Teaching Assistant responsible for providing detailed, actionable feedback to students on their writing assignments."
        "\n\n"
        "Instructions: The user will provide the student's assignment text and high-level feedback, both enclosed in triple backticks. Your task is to identify specific strengths in the student's work related to the feedback."
        "\n\n"
        "For each identified strength, include the following elements:"
        "1. **strength**: A concise description of the positive aspect."
        "2. **original_text**: The exact excerpt from the student's assignment where the strength is demonstrated (must match verbatim)."
        "3. **criterion_from_rubric**: The relevant grading criterion. If not explicitly provided, infer an appropriate criterion based on the feedback."
        "4. **comments**: Additional reinforcement, clarification, or advice."
        "\n\n"
        "\n\n"
        "Ensure that the `original_text` aligns with an exact excerpt from the student's assignment. Do not paraphrase or generalize this text."
    )

    user_prompt = ("Student assignment: ```"
        "{}\n"
        "```\n\n"
        "High level feedback: ```"
        "{}\n"
        "```").format(student_assignment, grading_feedback)

    improvements = get_feeback(system_prompt_improvements, user_prompt, schemas.response_format_improvements)['improvements']
    strengths = get_feeback(system_prompt_strengths, user_prompt, schemas.response_format_strengths)['strengths']

    return improvements, strengths

def create_feedback_html_with_css(student_assignment, improvements, strengths):
    """
    Displays specifics strengths and areas for improvement in the original essay text with hover effects.
    Highlights the revised text in the essay and provides a tooltip with the details of the improvement.
    """

    css = """
    <style>
    .feedback-tooltip {
        position: relative;
        display: inline-block;
        cursor: pointer;
        background-color: yellow;
        margin-bottom: 10px;  /* Adds space between highlighted sections */
        padding: 5px;  /* Adds padding for better visibility */
        border: 1px solid black;  /* Adds a border to the highlighted section */
        border-radius: 4px;  /* Adds rounded corners */
    }
    .feedback-tooltip .tooltiptext {
        visibility: hidden;
        width: 300px;
        background-color: white;
        color: black;
        text-align: left;
        border: 1px solid #ddd;
        padding: 10px;
        border-radius: 6px;
        box-shadow: 0 0 10px rgba(0, 0, 0, 0.1);
        position: absolute;
        z-index: 1;
        top: 100%;
        left: 50%;
        margin-left: -150px;
    }
    .feedback-tooltip:hover .tooltiptext {
        visibility: visible;
    }
    </style>
    """

    html_content = student_assignment

    for improvement in improvements:
        original = improvement.get('original_text', '')
        improvement_text = improvement['improvement']
        reason_for_suggestion = improvement['reason_for_suggestion']
        criterion = improvement['criterion_from_rubric']
        examples = improvement['examples']


        if original in student_assignment:
            tooltip_html = f"""
            <div class="feedback-tooltip">
                {original}
                <span class="tooltiptext">
                    <strong>Criterion:</strong> {criterion}<br>
                    <strong>Area for Improvement:</strong> {improvement_text}<br>
                    <strong>Reason:</strong> {reason_for_suggestion}<br>
                    <strong>Example:</strong> {examples}
                </span>
            </div>
            """
            html_content = html_content.replace(original, tooltip_html)
        else:
            # TODO: When original in not the student assignment: use a flash LLM to disambiguate (e.g. refers to a missing element), and identify correct position (such as missing title should be at the top)
            tooltip_html = f"""
            <div class="feedback-tooltip">
                {" "}
                <span class="tooltiptext">
                    <strong>Criterion:</strong> {criterion}<br>
                    <strong>Area for Improvement:</strong> {original}<br>
                    <strong>Reason:</strong> {reason_for_suggestion}<br>
                    <strong>Example:</strong> {examples}
                </span>
            </div>
            """

            html_content = tooltip_html + html_content

    for strength in strengths:
        original = strength.get('original_text', '')
        strength_text = strength['strength']
        criterion = strength['criterion_from_rubric']
        comments = strength['comments']
        if original in student_assignment:
            tooltip_html = f"""
            <div class="improvement-tooltip">
                {original}
                <span class="tooltiptext">
                    <strong>Criterion:</strong> {criterion}<br>
                    <strong>Strength:</strong> {strength_text}<br>
                    <p>{comments}</p>
                </span>
            </div>
            """
            html_content = html_content.replace(original, tooltip_html)

        else:
            # TODO: When original in not the student assignment: use a flash LLM to disambiguate: see improvements loop for similar situation.

            tooltip_html = f"""
            <div class="feedback-tooltip">
                {" "}
                <span class="tooltiptext">
                    <strong>Criterion:</strong> {criterion}<br>
                    <strong>Strength:</strong> {original}<br>
                    <p>{comments}</p>
                </span>
            </div>
            """
            html_content = tooltip_html + html_content
    return css + html_content
